import os
import re
import zipfile
import hashlib
import logging
from typing import Dict, Any, Optional
from xml.etree import ElementTree

logger = logging.getLogger(__name__)

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    logger.warning("PyPDF2 o'rnatilmagan — PDF tahlil cheklangan bo'ladi")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx o'rnatilmagan — DOCX tahlil cheklangan bo'ladi")

from config import SUSPICIOUS_PERMISSIONS, SUSPICIOUS_VIDEO_PATTERNS


def calculate_file_hash(file_path: str) -> str:
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.2f} MB"


def analyze_pdf(file_path: str) -> Dict[str, Any]:
    result = {
        "type": "PDF",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)

        if file_size > 50 * 1024 * 1024:
            result["score"] -= 10
            result["warnings"].append(
                f"⚠️ Katta hajm: {result['file_size']} (shubhali)"
            )
        else:
            result["details"].append(f"✅ Hajm: {result['file_size']}")

        if not HAS_PYPDF2:
            result["warnings"].append("⚠️ PyPDF2 o'rnatilmagan — chuqur tahlil cheklangan")
            return result

        with open(file_path, "rb") as f:
            try:
                reader = PyPDF2.PdfReader(f)
            except Exception as e:
                result["score"] -= 20
                result["warnings"].append(f"⚠️ PDF o'qishda xato: {str(e)[:80]}")
                return result

            if reader.is_encrypted:
                result["score"] -= 5
                result["warnings"].append("🔒 Parol himoyalangan PDF")
            else:
                result["details"].append("✅ Parol: Yo'q (ochiq)")

            meta = reader.metadata
            if meta:
                if meta.author:
                    result["metadata"]["Muallif"] = str(meta.author)
                if meta.creator:
                    result["metadata"]["Yaratuvchi"] = str(meta.creator)
                if meta.producer:
                    result["metadata"]["Ishlab chiqaruvchi"] = str(meta.producer)
                if meta.creation_date:
                    result["metadata"]["Yaratilgan sana"] = str(meta.creation_date)
                if meta.modification_date:
                    result["metadata"]["O'zgartirilgan sana"] = str(meta.modification_date)

                if result["metadata"]:
                    result["details"].append("✅ Metadata: Mavjud")
                else:
                    result["warnings"].append("⚠️ Metadata: Bo'sh (shubhali)")
                    result["score"] -= 5

            pages = len(reader.pages)
            result["details"].append(f"📄 Sahifalar: {pages}")

            js_found = False
            links_found = 0
            suspicious_links = []

            for page in reader.pages:
                if "/Annots" in page:
                    annotations = page["/Annots"]
                    if annotations:
                        for annot in annotations:
                            try:
                                annot_obj = annot.get_object()
                                if "/A" in annot_obj:
                                    action = annot_obj["/A"]
                                    if "/URI" in action:
                                        links_found += 1
                                        uri = str(action["/URI"])
                                        if any(kw in uri.lower() for kw in
                                               ["login", "verify", "password", "account"]):
                                            suspicious_links.append(uri[:80])
                                    if "/JS" in action or "/JavaScript" in action:
                                        js_found = True
                            except Exception:
                                pass

            if js_found:
                result["score"] -= 30
                result["warnings"].append("🔴 JavaScript aniqlandi: Zararli kod xavfi")
            else:
                result["details"].append("✅ JavaScript: Topilmadi")

            if suspicious_links:
                result["score"] -= 15
                result["warnings"].append(
                    f"⚠️ Shubhali havolalar: {len(suspicious_links)} ta"
                )
                for link in suspicious_links[:3]:
                    result["warnings"].append(f"   └─ {link}")
            elif links_found > 0:
                result["details"].append(f"🔗 Havolalar: {links_found} ta")
            else:
                result["details"].append("✅ Havolalar: Yo'q")

    except Exception as e:
        logger.error("PDF tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def analyze_docx(file_path: str) -> Dict[str, Any]:
    result = {
        "type": "DOCX",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)
        result["details"].append(f"✅ Hajm: {result['file_size']}")

        has_macros = False
        external_links = []
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                file_list = z.namelist()

                for name in file_list:
                    if "vbaProject" in name or name.endswith(".bin"):
                        has_macros = True
                        break

                for name in file_list:
                    if name.endswith(".rels"):
                        try:
                            with z.open(name) as rels_file:
                                content = rels_file.read().decode("utf-8", errors="ignore")
                                ext_links = re.findall(
                                    r'Target="(https?://[^"]+)"', content
                                )
                                external_links.extend(ext_links)
                        except Exception:
                            pass
        except zipfile.BadZipFile:
            result["score"] -= 15
            result["warnings"].append("⚠️ ZIP formati buzilgan — shubhali fayl")

        if has_macros:
            result["score"] -= 30
            result["warnings"].append("🔴 VBA Makros aniqlandi: Zararli kod xavfi yuqori!")
        else:
            result["details"].append("✅ Makros: Topilmadi")

        if external_links:
            result["score"] -= 10
            result["warnings"].append(
                f"⚠️ Tashqi havolalar: {len(external_links)} ta"
            )
            for link in external_links[:3]:
                result["warnings"].append(f"   └─ {link[:80]}")
        else:
            result["details"].append("✅ Tashqi havolalar: Yo'q")

        if HAS_DOCX:
            try:
                doc = docx.Document(file_path)

                props = doc.core_properties
                if props.author:
                    result["metadata"]["Muallif"] = str(props.author)
                if props.last_modified_by:
                    result["metadata"]["Oxirgi tahrir"] = str(props.last_modified_by)
                if props.created:
                    result["metadata"]["Yaratilgan"] = str(props.created)
                if props.modified:
                    result["metadata"]["O'zgartirilgan"] = str(props.modified)
                if props.revision:
                    result["metadata"]["Tahrir soni"] = str(props.revision)

                if result["metadata"]:
                    result["details"].append("✅ Metadata: Mavjud")

                embedded_count = 0
                for rel in doc.part.rels.values():
                    if "oleObject" in str(rel.reltype) or "embedding" in str(rel.reltype):
                        embedded_count += 1

                if embedded_count > 0:
                    result["score"] -= 15
                    result["warnings"].append(
                        f"⚠️ Embedded obyektlar: {embedded_count} ta (xavfli bo'lishi mumkin)"
                    )
                else:
                    result["details"].append("✅ Embedded obyektlar: Yo'q")

                hidden_text_found = False
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run.font.hidden:
                            hidden_text_found = True
                            break
                    if hidden_text_found:
                        break

                if hidden_text_found:
                    result["score"] -= 10
                    result["warnings"].append("⚠️ Yashirin matn aniqlandi")
                else:
                    result["details"].append("✅ Yashirin matn: Yo'q")

                result["details"].append(f"📝 Paragraflar: {len(doc.paragraphs)}")
                result["details"].append(f"📊 Jadvallar: {len(doc.tables)}")

            except Exception as e:
                result["warnings"].append(f"⚠️ DOCX chuqur tahlil xatosi: {str(e)[:60]}")
        else:
            result["warnings"].append("⚠️ python-docx o'rnatilmagan — chuqur tahlil cheklangan")

    except Exception as e:
        logger.error("DOCX tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def analyze_apk(file_path: str) -> Dict[str, Any]:
    result = {
        "type": "APK",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
        "permissions": [],
        "suspicious_permissions": [],
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)
        result["details"].append(f"✅ Hajm: {result['file_size']}")

        if file_size < 100 * 1024:
            result["score"] -= 15
            result["warnings"].append("⚠️ Juda kichik APK — dropper yoki stub bo'lishi mumkin")

        try:
            with zipfile.ZipFile(file_path, "r") as apk_zip:
                file_list = apk_zip.namelist()

                has_manifest = "AndroidManifest.xml" in file_list
                has_dex = any(f.endswith(".dex") for f in file_list)
                has_resources = "resources.arsc" in file_list

                if has_manifest:
                    result["details"].append("✅ AndroidManifest.xml: Mavjud")
                else:
                    result["score"] -= 20
                    result["warnings"].append("🔴 AndroidManifest.xml: Topilmadi (noto'g'ri APK)")

                if has_dex:
                    dex_files = [f for f in file_list if f.endswith(".dex")]
                    result["details"].append(f"✅ DEX fayllari: {len(dex_files)} ta")
                    if len(dex_files) > 5:
                        result["score"] -= 5
                        result["warnings"].append(
                            f"⚠️ Ko'p DEX: {len(dex_files)} ta (multidex yoki obfuscation)"
                        )
                else:
                    result["score"] -= 20
                    result["warnings"].append("🔴 DEX fayl topilmadi")

                if has_resources:
                    result["details"].append("✅ Resources: Mavjud")

                has_signature = any(
                    f.startswith("META-INF/") and (
                        f.endswith(".RSA") or f.endswith(".DSA") or f.endswith(".SF")
                    )
                    for f in file_list
                )
                if has_signature:
                    result["details"].append("✅ Imzo (signature): Mavjud")
                else:
                    result["score"] -= 15
                    result["warnings"].append("⚠️ Imzo topilmadi — rasmiy bo'lmagan APK")

                if has_manifest:
                    try:
                        manifest_data = apk_zip.read("AndroidManifest.xml")
                        manifest_text = manifest_data.decode("utf-8", errors="ignore")

                        all_permissions = []
                        for perm in SUSPICIOUS_PERMISSIONS:
                            if perm in manifest_text or perm.lower() in manifest_text.lower():
                                all_permissions.append(perm)

                        result["permissions"] = all_permissions

                        suspicious = [
                            p for p in all_permissions
                            if p in SUSPICIOUS_PERMISSIONS
                        ]
                        result["suspicious_permissions"] = suspicious

                        if suspicious:
                            penalty = min(len(suspicious) * 5, 30)
                            result["score"] -= penalty
                            result["warnings"].append(
                                f"⚠️ Shubhali ruxsatlar ({len(suspicious)} ta):"
                            )
                            for p in suspicious:
                                result["warnings"].append(f"   └─ {p}")
                        else:
                            result["details"].append("✅ Shubhali ruxsatlar: Topilmadi")

                    except Exception as e:
                        result["warnings"].append(
                            f"⚠️ Manifest o'qishda xato: {str(e)[:60]}"
                        )

                so_files = [f for f in file_list if f.endswith(".so")]
                if so_files:
                    result["details"].append(f"📦 Native kutubxonalar: {len(so_files)} ta")

        except zipfile.BadZipFile:
            result["score"] -= 30
            result["warnings"].append("🔴 ZIP formati buzilgan — noto'g'ri yoki buzilgan APK")

    except Exception as e:
        logger.error("APK tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def analyze_video(file_path: str, file_name: str = "") -> Dict[str, Any]:
    result = {
        "type": "Video",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)

        ext = os.path.splitext(file_path)[1].lower()
        if not ext and file_name:
            ext = os.path.splitext(file_name)[1].lower()

        result["metadata"]["Format"] = ext.lstrip(".").upper() if ext else "Noma'lum"
        result["details"].append(f"✅ Hajm: {result['file_size']}")
        result["details"].append(f"🎬 Format: {result['metadata']['Format']}")

        if file_size < 10 * 1024:
            result["score"] -= 20
            result["warnings"].append("⚠️ Juda kichik video — yashirin fayl bo'lishi mumkin")

        name_lower = (file_name or os.path.basename(file_path)).lower()
        found_patterns = []
        for pattern in SUSPICIOUS_VIDEO_PATTERNS:
            if pattern in name_lower:
                found_patterns.append(pattern)

        if found_patterns:
            result["score"] -= len(found_patterns) * 8
            result["warnings"].append(
                f"⚠️ Shubhali nom: {', '.join(found_patterns)}"
            )
        else:
            result["details"].append("✅ Fayl nomi: Normal")

        name_parts = file_name.split(".") if file_name else []
        if len(name_parts) > 2:
            suspicious_exts = [".exe", ".bat", ".cmd", ".scr", ".vbs", ".js", ".ps1"]
            for p in name_parts[1:]:
                if f".{p.lower()}" in suspicious_exts:
                    result["score"] -= 30
                    result["warnings"].append(
                        f"🔴 Ikkilangan kengaytma: .{p} (zararli fayl yashirilgan!)"
                    )

        try:
            with open(file_path, "rb") as f:
                header = f.read(12)

            is_real_video = False
            if b"ftyp" in header[:12]:
                is_real_video = True
            elif header[:4] == b"RIFF" and header[8:12] == b"AVI ":
                is_real_video = True
            elif header[:4] == b"\x1a\x45\xdf\xa3":
                is_real_video = True
            elif header[:3] == b"FLV":
                is_real_video = True

            if is_real_video:
                result["details"].append("✅ Video header: Haqiqiy video fayl")
            else:
                result["score"] -= 25
                result["warnings"].append(
                    "🔴 Video header mos kelmaydi — boshqa turdagi fayl yashirilgan bo'lishi mumkin"
                )

        except Exception:
            result["warnings"].append("⚠️ Header tekshiruvda xato")

        if file_size > 500 * 1024 * 1024:
            result["warnings"].append("⚠️ Juda katta fayl — steganografiya ehtimoli bor")
            result["score"] -= 5

    except Exception as e:
        logger.error("Video tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def analyze_image(file_path: str, file_name: str = "") -> Dict[str, Any]:
    result = {
        "type": "Image",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)

        ext = os.path.splitext(file_path)[1].lower()
        if not ext and file_name:
            ext = os.path.splitext(file_name)[1].lower()

        result["metadata"]["Format"] = ext.lstrip(".").upper() if ext else "Noma'lum"
        result["details"].append(f"✅ Hajm: {result['file_size']}")

        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            with Image.open(file_path) as img:
                result["details"].append(f"🖼️ O'lcham: {img.width}x{img.height}")
                result["metadata"]["Asl Format"] = img.format

                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    has_gps = False
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        if tag == "GPSInfo":
                            has_gps = True
                        elif tag in ("Make", "Model", "DateTimeOriginal", "Software", "UserComment"):
                            result["metadata"][str(tag)] = str(value)[:50]
                    
                    if has_gps:
                        result["warnings"].append("⚠️ EXIF: GPS joylashuv ma'lumoti mavjud")
                        result["score"] -= 10
                    else:
                        result["details"].append("✅ EXIF: GPS topilmadi")
        except ImportError:
            result["warnings"].append("⚠️ Pillow kutubxonasi yo'q — EXIF tekshirilmaydi")
        except Exception as e:
            result["warnings"].append(f"⚠️ Rasm tahlil xatosi: {str(e)[:40]}")

        # EOF tekshiruvi (Steganography / Yashirin ma'lumot)
        try:
            with open(file_path, "rb") as f:
                content = f.read()
                # Keng tarqalgan EOF markerlar
                eof_markers = {
                    '.jpg': b'\xff\xd9',
                    '.jpeg': b'\xff\xd9',
                    '.png': b'\x49\x45\x4e\x44\xae\x42\x60\x82',
                    '.gif': b'\x3b',
                }
                
                marker = eof_markers.get(ext)
                if marker:
                    idx = content.rfind(marker)
                    if idx != -1:
                        trailing_data = content[idx + len(marker):]
                        if len(trailing_data) > 100:
                            result["warnings"].append(
                                f"🔴 Yashirin ma'lumot: EOF dan keyin {len(trailing_data)} bayt ortiqcha ma'lumot bor!"
                            )
                            result["score"] -= 30
        except Exception:
            pass

    except Exception as e:
        logger.error("Rasm tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def analyze_archive(file_path: str, file_name: str = "") -> Dict[str, Any]:
    result = {
        "type": "Archive",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)
        result["details"].append(f"✅ Hajm: {result['file_size']}")

        ext = os.path.splitext(file_name)[1].lower() if file_name else ".zip"

        if ext == '.zip':
            try:
                with zipfile.ZipFile(file_path, 'r') as z:
                    file_list = z.namelist()
                    num_files = len(file_list)
                    result["details"].append(f"📦 Jami fayllar: {num_files} ta")

                    uncompressed_size = sum(info.file_size for info in z.infolist())
                    
                    # Zip bomb tekshiruvi (compression ratio > 100x va hajm > 1GB)
                    if file_size > 0 and uncompressed_size / file_size > 100 and uncompressed_size > 10**9:
                        result["warnings"].append("🔴 ZIP Bomb aniqlandi (juda yuqori siqish nisbati)!")
                        result["score"] -= 40

                    executables = [f for f in file_list if f.lower().endswith(('.exe', '.bat', '.cmd', '.scr', '.vbs', '.js', '.ps1', '.hta'))]
                    if executables:
                        result["warnings"].append(f"🔴 Arxivda ishga tushuvchi fayllar bor: {len(executables)} ta")
                        result["score"] -= 25
                        for ex in executables[:3]:
                            result["warnings"].append(f"   └─ {ex}")

                    if any(info.flag_bits & 0x1 for info in z.infolist()):
                        result["warnings"].append("🔒 Arxiv parol himoyalangan")
                        result["score"] -= 5

            except zipfile.BadZipFile:
                result["warnings"].append("🔴 ZIP formati buzilgan")
                result["score"] -= 20
        else:
            result["warnings"].append(f"⚠️ {ext} formati chuqur tahlil qilinmadi (faqat ZIP)")

    except Exception as e:
        logger.error("Arxiv tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def analyze_js(file_path: str, file_name: str = "") -> Dict[str, Any]:
    result = {
        "type": "JavaScript",
        "score": 100,
        "details": [],
        "warnings": [],
        "metadata": {},
        "file_hash": "",
        "file_size": "",
    }

    try:
        file_size = os.path.getsize(file_path)
        result["file_size"] = format_size(file_size)
        result["file_hash"] = calculate_file_hash(file_path)
        result["details"].append(f"✅ Hajm: {result['file_size']}")

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            lines = content.count('\n') + 1
            result["details"].append(f"📝 Qatorlar soni: {lines}")

            lower_content = content.lower()

            # Xavfli funksiyalar
            dangerous_funcs = ['eval(', 'document.write(', 'setTimeout(', 'setInterval(', 'Function(']
            found_funcs = [func for func in dangerous_funcs if func in lower_content]
            if found_funcs:
                result["warnings"].append(f"⚠️ Xavfli funksiyalar: {', '.join(found_funcs[:3])}")
                result["score"] -= len(found_funcs) * 5

            # Obfuscation tekshiruvi (uzun qatorlar va g'alati belgilar)
            if lines < 3 and len(content) > 1000:
                result["warnings"].append("🔴 Kod siqilgan (Minified) yoki Obfuskatsiya qilingan")
                result["score"] -= 15

            if 'atob(' in lower_content or 'btoa(' in lower_content:
                result["warnings"].append("⚠️ Base64 kodlash/dekodlash aniqlandi")
                result["score"] -= 10

            hex_escapes = len(re.findall(r'\\x[0-9a-fA-F]{2}', content))
            if hex_escapes > 20:
                result["warnings"].append(f"🔴 Obfuskatsiya belgilari (hex escapes): {hex_escapes} ta")
                result["score"] -= 20

            # Tashqi so'rovlar
            if 'xmlhttprequest' in lower_content or 'fetch(' in lower_content or '$.ajax(' in lower_content:
                result["warnings"].append("⚠️ Kod tashqi serverga so'rov yuboradi (AJAX/Fetch)")
                result["score"] -= 5

        except Exception as e:
            result["warnings"].append(f"⚠️ JavaScript o'qishda xato: {str(e)[:40]}")

    except Exception as e:
        logger.error("JS tahlil xatosi: %s", e)
        result["score"] -= 20
        result["warnings"].append(f"❌ Tahlil xatosi: {str(e)[:80]}")

    result["score"] = max(0, result["score"])
    return result


def detect_file_type(file_name: str) -> Optional[str]:
    ext = os.path.splitext(file_name)[1].lower()

    from config import SUPPORTED_EXTENSIONS
    for file_type, extensions in SUPPORTED_EXTENSIONS.items():
        if ext in extensions:
            return file_type

    return None
